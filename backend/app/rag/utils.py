import re
from io import BytesIO
import markdown
import hashlib
from typing import List

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings


def extract_text_from_pdf(file_obj: BytesIO) -> str:
    from pdfminer.high_level import extract_text
    return extract_text(file_obj)
    

def extract_text_from_docx(file_obj: BytesIO) -> str:
    from docx import Document
    doc = Document(file_obj)
    return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

# def extract_text_from_md_html(file_path: str) -> str:
#     """Extracts text from a Markdown or HTML file."""
#     with open(file_path, "r", encoding="utf-8") as f:
#         md_content = f.read()
    
#     html_content = markdown.markdown(md_content)
#     # Clean up the HTML by removing unnecessary tags (keeping only text)
#     text = re.sub(r'<[^>]+>', '', html_content) 
    
#     # Clean up extra spaces and join text
#     text = ' '.join(text.split())

#     return text

def extract_text_from_md_html(file_obj: BytesIO) -> str:
    from bs4 import BeautifulSoup
    content = file_obj.read().decode('utf-8')
    if content.startswith('<!DOCTYPE') or content.startswith('<html'):
        soup = BeautifulSoup(content, 'html.parser')
        return soup.get_text()
    return content 
    
def clean_text(text: str) -> str:
    """Preprocess and clean extracted text while preserving currency symbols and percent signs.."""
    if not text:
        return ""

    text = " ".join(text.split())
    text = re.sub(r"[^a-zA-Z0-9\s$€£¥₹%:/.-]", "", text)
    text = text.strip().lower()  # Normalize text
    return text

def split_document(text, filename : str, file_type: str, user_id: str):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size = 300,
        chunk_overlap = 50,
        separators= ["\n"," "]
    )

    chunks = text_splitter.split_text(text)

    documents = []
    for i, chunk in enumerate(chunks):
        chunk_hash = hashlib.md5(chunk.encode()).hexdigest()
        metadata = {
            "document_name": filename,  
            "user_id": user_id,              
            "chunk_id": chunk_hash,  
            "chunk_index": i,                 
            "content_type": file_type, 
            "content" : chunk    
        }

        doc = Document(page_content = chunk, metadata = metadata)
        documents.append(doc)
    
    return documents
    

def get_embedding(documents: List[Document]) -> List[dict]:
    embeddings = []
    model = OpenAIEmbeddings(model="text-embedding-3-small")
    for document in documents:
        vector = model.embed_documents([document.page_content])[0]
        embeddings.append({
            "id": document.metadata["chunk_id"],
            "values": vector,
            "metadata": document.metadata
        })
    return embeddings

def query_embedding(query: str):
    model = OpenAIEmbeddings(model="text-embedding-3-small")
    query_embedding = {
        "values": model.embed_query(query)
    }
        
    return query_embedding
