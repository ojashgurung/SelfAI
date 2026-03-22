import os
from io import BytesIO
from typing import Optional
from fastapi import HTTPException

from graph.ingestors.base import IngestPayload, IngestedDocument
from graph.utils.hash import sha256_text


class DocumentIngestor:
    platform = "documents"

    def ingest(
        self,
        *,
        file_bytes: bytes,
        filename: str,
        content_type: Optional[str] = None,
        user_id: str,
    ) -> IngestPayload:
        ext = os.path.splitext(filename)[1].lower()
        text = self._extract_text(file_bytes, ext)

        if not text or not text.strip():
            raise HTTPException(
                status_code=400,
                detail="Could not extract text from file"
            )

        content_hash = sha256_text(text)
        external_id = f"upload:file:{content_hash}"

        return IngestPayload(
            platform="documents",
            account_id=user_id,
            display_name="Documents",
            source_metadata={
                "filename": filename,
                "file_size": len(file_bytes),
                "content_type": content_type,
            },
            documents=[
                IngestedDocument(
                    external_id=external_id,
                    doc_type=ext.lstrip("."),
                    title=filename,
                    text=text,
                    url=None,
                    updated_at_source=None,
                    metadata={
                        "filename": filename,
                        "file_size": len(file_bytes),
                        "content_type": content_type,
                    },
                )
            ],
        )

    def _extract_text(self, file_bytes: bytes, ext: str) -> str:
        buffer = BytesIO(file_bytes)

        if ext == ".pdf":
            from pdfminer.high_level import extract_text
            return extract_text(buffer)

        elif ext == ".docx":
            from docx import Document
            doc = Document(buffer)
            return "\n".join([p.text for p in doc.paragraphs])

        elif ext in (".md", ".html", ".txt"):
            return buffer.read().decode("utf-8", errors="ignore")

        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {ext}"
        )